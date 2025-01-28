from docx import Document

realistic_website_data = [
    (1, "https://www.google.com", "Search Engine", "Minimalistic design; clean interface", "Lack of customization options", "Essential tool for information retrieval"),
    (2, "https://www.youtube.com", "Video Sharing Platform", "Intuitive layout; easy navigation", "Too many ads disrupt user experience", "Excellent for entertainment and learning"),
    (3, "https://www.facebook.com", "Social Networking", "Organized layout; responsive design", "Overwhelming clutter in some sections", "Great for staying connected but requires cautious use"),
    (4, "https://www.instagram.com", "Photo and Video Sharing Platform", "Visually appealing; smooth transitions", "Repetitive design elements", "Popular platform for visual content sharing"),
    (5, "https://www.twitter.com", "Microblogging and Social Networking", "Simple and effective layout", "Font size and color contrasts can be improved", "Effective for quick information and trends"),
    (6, "https://www.linkedin.com", "Professional Networking", "Professional-looking interface; well-organized", "Occasional slow loading of profiles", "Valuable for career development and networking"),
    (7, "https://www.wikipedia.org", "Online Encyclopedia", "Straightforward layout; easy-to-read typography", "Outdated design aesthetics", "Reliable starting point for research"),
    (8, "https://www.amazon.com", "E-commerce Platform", "User-friendly categorization; good use of white space", "Too many elements on product pages", "Convenient for online shopping"),
    (9, "https://www.netflix.com", "Streaming Service", "Sleek and modern interface; dark mode", "Limited customization of homepage layout", "Leading platform for diverse entertainment"),
    (10, "https://www.reddit.com", "Community Discussion Forum", "Minimalist design for discussions", "Complex for new users; lacks visual appeal", "Engaging platform for discussions and information"),
    (11, "https://www.pinterest.com", "Image Sharing and Social Media", "Grid-based layout; visually pleasing", "Slow loading of high-resolution images", "Great for discovering and saving ideas"),
    (12, "https://www.tiktok.com", "Short-form Video Sharing", "Dynamic and interactive design", "Cluttered homepage with too many recommendations", "Fun platform for short videos"),
    (13, "https://www.ebay.com", "Online Marketplace", "Clear categorization; easy checkout process", "Outdated and inconsistent design", "Useful for buying and selling a variety of items"),
    (14, "https://www.bing.com", "Search Engine", "Visually appealing homepage images", "Search results page feels cluttered", "Viable alternative search engine"),
    (15, "https://www.yahoo.com", "Web Portal and Search Engine", "Comprehensive dashboard layout", "Overcrowded homepage with ads", "Comprehensive portal but needs modernization"),
    (16, "https://www.whatsapp.com", "Messaging Application", "Clean and straightforward interface", "Limited UI customization options", "Reliable for instant messaging"),
    (17, "https://www.twitch.tv", "Live Streaming Platform", "Interactive and engaging layout", "Busy homepage with too much information", "Leading platform for live gaming and creative content"),
    (18, "https://www.nytimes.com", "News Publication", "Elegant typography; organized content", "Too many ads disrupt reading flow", "Trusted source for news and analysis"),
    (19, "https://www.cnn.com", "News and Media Outlet", "Multimedia integration; easy navigation", "Too many pop-ups and ads", "Comprehensive news source but requires critical consumption"),
    (20, "https://www.bbc.com", "News and Media Organization", "Clean and professional layout", "Limited interactive elements", "Respected source for international news"),
    (21, "https://www.quora.com", "Question-and-Answer Platform", "Simple layout; easy to follow threads", "Lack of visual appeal", "Informative platform for knowledge sharing"),
    (22, "https://www.medium.com", "Online Publishing Platform", "Clean and distraction-free reading interface", "Limited font options for writers", "Great for in-depth articles and diverse perspectives"),
    (23, "https://www.spotify.com", "Music Streaming Service", "Dark mode; seamless navigation", "Limited customization of playlist views", "Excellent for music discovery and listening"),
    (24, "https://www.apple.com", "Technology Company Website", "Sleek and innovative design; intuitive navigation", "Over-reliance on images; slow loading on poor networks", "Informative site for Apple products and services"),
    (25, "https://www.microsoft.com", "Technology Company Website", "Professional look; consistent design language", "Some sections feel cluttered", "Resourceful for Microsoft products and services"),
    (26, "https://www.adobe.com", "Software Company Website", "Visually stunning design; good use of color schemes", "Complex navigation for new users", "Essential for creative professionals"),
    (27, "https://www.canva.com", "Graphic Design Platform", "User-friendly drag-and-drop interface", "Occasional lag when editing designs", "Ideal for quick and easy graphic design"),
    (28, "https://www.salesforce.com", "Customer Relationship Management", "Modern design; customizable dashboards", "Overwhelming for first-time users", "Powerful solution for enterprise customer management"),
    (29, "https://www.shopify.com", "E-commerce Platform", "Clean templates; intuitive design for sellers", "Limited advanced customization options", "Great for small to medium-sized online stores"),
    (30, "https://www.airbnb.com", "Online Lodging Marketplace", "Visually attractive listings; intuitive search", "Inconsistent loading speeds for images", "Excellent for finding diverse lodging options"),
]
# Create a new document with the filled table
doc = Document()
doc.add_heading("Website Evaluation Report", level=1)

# Add table headers
headers = ["Srno", "WebSiteURL", "Purpose Of Website", "Things Liked In Website", 
           "Things Disliked In Website", "Overall Evaluation"]
table = doc.add_table(rows=1, cols=len(headers))
header_row = table.rows[0]
for idx, header in enumerate(headers):
    header_row.cells[idx].text = header

# Fill the table with the actual data
for entry in realistic_website_data:
    row = table.add_row()
    for idx, value in enumerate(entry):
        row.cells[idx].text = str(value)

# Save the file
file_path = "./Assg_1.docx"
doc.save(file_path)